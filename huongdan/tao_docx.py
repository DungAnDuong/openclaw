"""
Tao file HuongDan_CaiDat_OpenClaw.docx bang python-docx
Chay: python huongdan/tao_docx.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "HuongDan_CaiDat_OpenClaw.docx")

doc = Document()

# --- Thiet lap trang ---
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x10, 0x7A, 0x40)
    return p


def para(text, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD6, 0x36, 0x18)
    # Background
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F3F4F6')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shading)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    return p


def numbered(text):
    return doc.add_paragraph(text, style='List Number')


def table_row(tbl, cells):
    row = tbl.add_row()
    for i, c in enumerate(cells):
        row.cells[i].text = c
    return row


def separator():
    doc.add_paragraph("─" * 60)


# ============================================================
# TRANG BIA
# ============================================================
doc.add_heading("", 0)
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("HƯỚNG DẪN CÀI ĐẶT VÀ VẬN HÀNH")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run("OpenClaw — Trợ Lý AI Cá Nhân")
run2.bold = True
run2.font.size = Pt(18)

doc.add_paragraph()
ver_para = doc.add_paragraph()
ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_para.add_run("Phiên bản: 2026.4.24 | Cập nhật: Tháng 6/2026")

doc.add_paragraph()
doc.add_paragraph()

# ============================================================
# MUC LUC
# ============================================================
heading("MỤC LỤC", 1)
toc_items = [
    "1. Giới thiệu tổng quan",
    "2. Yêu cầu hệ thống",
    "3. Hướng dẫn cài đặt chi tiết",
    "   3.1 Cài đặt Node.js",
    "   3.2 Cài đặt Git",
    "   3.3 Tải source code",
    "   3.4 Cài đặt OpenClaw",
    "   3.5 Cấu hình môi trường",
    "   3.6 Chạy cài đặt tự động",
    "   3.7 Khởi động chương trình",
    "   3.8 Kiểm tra hoạt động",
    "4. Hướng dẫn sử dụng file cấu hình",
    "   4.1 llm_config.json",
    "   4.2 bot_config.json",
    "   4.3 app_config.json",
    "   4.4 paths.json",
    "5. Giải thích các lệnh sử dụng",
    "6. Lỗi thường gặp và cách khắc phục",
    "7. Hướng dẫn cập nhật phiên bản",
    "8. Backup và Restore cấu hình",
    "9. Sơ đồ hoạt động hệ thống",
]
for item in toc_items:
    bullet(item)

doc.add_page_break()

# ============================================================
# 1. GIOI THIEU
# ============================================================
heading("1. GIỚI THIỆU TỔNG QUAN", 1)

para(
    "OpenClaw là trợ lý AI cá nhân mạnh mẽ, chạy trên thiết bị của bạn và kết nối "
    "với các ứng dụng nhắn tin phổ biến. Không giống các dịch vụ AI đám mây thông "
    "thường, OpenClaw chạy tại chỗ (local-first) giúp bảo mật dữ liệu cá nhân."
)

heading("Tính năng chính:", 2)
features = [
    "Hỗ trợ đa kênh: WhatsApp, Telegram, Discord, Slack, Microsoft Teams, Signal, "
    "Zalo, Mattermost, Matrix, LINE, Feishu, Tlon, Twitch, và nhiều hơn nữa",
    "Hỗ trợ nhiều nhà cung cấp AI: Anthropic Claude, OpenAI GPT, Google Gemini, "
    "Mistral, Groq, và các model tự host (Ollama, LM Studio)",
    "Giao diện quản lý web tại http://127.0.0.1:18789",
    "Tự động hóa bằng cron jobs và webhooks",
    "Hỗ trợ voice (giọng nói) trên macOS/iOS/Android",
    "Sandbox bảo mật — chạy agent trong môi trường cách ly",
    "Plugin ecosystem phong phú qua ClawHub",
]
for f in features:
    bullet(f)

heading("Kiến trúc hệ thống:", 2)
para(
    "OpenClaw hoạt động theo mô hình Gateway — một server trung tâm chạy trên máy "
    "của bạn, nhận tin nhắn từ các kênh chat, xử lý bằng AI, và trả về phản hồi. "
    "Cổng mặc định: 18789."
)

doc.add_page_break()

# ============================================================
# 2. YEU CAU HE THONG
# ============================================================
heading("2. YÊU CẦU HỆ THỐNG", 1)

heading("2.1 Hệ điều hành", 2)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
header_row = tbl.rows[0]
header_row.cells[0].text = "Hệ điều hành"
header_row.cells[1].text = "Phiên bản"
header_row.cells[2].text = "Ghi chú"
for cell in header_row.cells:
    cell.paragraphs[0].runs[0].bold = True

table_row(tbl, ["Windows 10", "Version 2004 (Build 19041)+", "Hỗ trợ đầy đủ"])
table_row(tbl, ["Windows 11", "Tất cả phiên bản", "Khuyến nghị"])
table_row(tbl, ["Windows Server", "2019/2022", "Hỗ trợ"])
doc.add_paragraph()

heading("2.2 Phần mềm yêu cầu", 2)
tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = 'Table Grid'
h2 = tbl2.rows[0]
for i, t in enumerate(["Phần mềm", "Phiên bản tối thiểu", "Khuyến nghị", "Bắt buộc"]):
    h2.cells[i].text = t
    h2.cells[i].paragraphs[0].runs[0].bold = True

table_row(tbl2, ["Node.js", "22.14.0", "24.x LTS", "CÓ"])
table_row(tbl2, ["npm", "10.x", "Đi kèm Node.js", "CÓ"])
table_row(tbl2, ["Git", "2.x", "Mới nhất", "Khuyến nghị"])
table_row(tbl2, ["pnpm", "9.x", "Tùy chọn", "Không"])
table_row(tbl2, ["Docker Desktop", "4.x", "Mới nhất", "Không (sandbox)"])
doc.add_paragraph()

heading("2.3 Phần cứng tối thiểu", 2)
hw_items = [
    "RAM: 4 GB (khuyến nghị 8 GB+)",
    "Đĩa cứng: 2 GB trống cho cài đặt cơ bản",
    "CPU: Bất kỳ x64 processor hiện đại",
    "Kết nối Internet: Bắt buộc để tải model và giao tiếp với API",
]
for h in hw_items:
    bullet(h)

heading("2.4 API Key (chọn ít nhất một)", 2)
api_items = [
    "Anthropic API Key — https://console.anthropic.com (model Claude)",
    "OpenAI API Key — https://platform.openai.com (model GPT)",
    "Google AI API Key — https://aistudio.google.com (model Gemini)",
    "HOẶC: Model tự host (Ollama/LM Studio) — không cần API key",
]
for a in api_items:
    bullet(a)

doc.add_page_break()

# ============================================================
# 3. HUONG DAN CAI DAT
# ============================================================
heading("3. HƯỚNG DẪN CÀI ĐẶT CHI TIẾT", 1)

heading("3.1 Cài đặt Node.js", 2)
para(
    "Node.js là nền tảng chạy OpenClaw. Bắt buộc phải cài đặt phiên bản 22.14.0 "
    "trở lên, khuyến nghị dùng phiên bản 24 LTS."
)

numbered("Truy cập: https://nodejs.org/en/download")
numbered("Chọn tab 'Prebuilt Installer' → Windows → x64")
numbered("Tải file .msi phiên bản LTS (24.x) và chạy installer")
numbered("Trong quá trình cài, chọn 'Add to PATH' (thường đã được chọn sẵn)")
numbered("Sau khi cài xong, mở Command Prompt (cmd) và kiểm tra:")

code_block("node --version\n# Kết quả mong đợi: v24.x.x\n\nnpm --version\n# Kết quả mong đợi: 10.x.x")

para("Nếu thấy lỗi 'node is not recognized', khởi động lại máy tính và thử lại.", color=(0xC0, 0x39, 0x2B))

heading("3.2 Cài đặt Git (khuyến nghị)", 2)
numbered("Truy cập: https://git-scm.com/download/win")
numbered("Tải file installer và chạy với cài đặt mặc định")
numbered("Kiểm tra sau khi cài:")
code_block("git --version\n# Kết quả mong đợi: git version 2.x.x")

heading("3.3 Tải source code (nếu phát triển từ source)", 2)
para("Nếu bạn chỉ muốn DÙNG OpenClaw (không phát triển), bỏ qua bước này và sang 3.4.")
code_block(
    "git clone https://github.com/openclaw/openclaw.git\n"
    "cd openclaw"
)

heading("3.4 Cài đặt OpenClaw", 2)
para("Cài đặt OpenClaw toàn cục (global) qua npm — đây là cách đơn giản nhất:", bold=False)

code_block(
    "# Cài đặt phiên bản mới nhất\n"
    "npm install -g openclaw@latest\n\n"
    "# Kiểm tra cài đặt thành công\n"
    "openclaw --version"
)

para("Hoặc dùng pnpm nếu đã cài:")
code_block("pnpm add -g openclaw@latest")

heading("3.5 Cấu hình môi trường", 2)
numbered("Mở thư mục dự án trong File Explorer")
numbered("Mở file config\\llm_config.json bằng Notepad hoặc VS Code")
numbered("Điền thông tin API key của bạn:")

code_block(
    '{\n'
    '  "provider": "anthropic",\n'
    '  "model": "claude-sonnet-4-6",\n'
    '  "api_key": "sk-ant-api03-YOUR-KEY-HERE",\n'
    '  "base_url": "",\n'
    '  "max_tokens": 8192,\n'
    '  "temperature": 1.0\n'
    '}'
)

numbered("Lưu file")
numbered("(Tùy chọn) Mở config\\app_config.json để thay đổi cổng nếu 18789 đã bị dùng")

heading("3.6 Chạy cài đặt tự động", 2)
para("Sau khi đã cấu hình xong file config, chạy install.bat:")
code_block(
    "# Nhấp đúp vào install.bat trong File Explorer\n"
    "# HOẶC mở cmd trong thư mục dự án và chạy:\n"
    "install.bat"
)

para("Script install.bat sẽ tự động:")
items = [
    "Kiểm tra Node.js, npm, Git",
    "Cài đặt openclaw từ npm",
    "Tạo các thư mục cần thiết (logs/, data/, backups/)",
    "Áp dụng cấu hình từ config/ vào ~/.openclaw/openclaw.json",
    "Chạy openclaw doctor để kiểm tra sức khỏe hệ thống",
]
for i in items:
    bullet(i)

heading("3.7 Khởi động chương trình", 2)
code_block(
    "# Nhấp đúp vào start.bat\n"
    "# HOẶC:\n"
    "start.bat"
)

para("Sau khi khởi động thành công, bạn sẽ thấy:")
code_block(
    "OpenClaw Gateway dang chay...\n"
    "Giao dien quan ly: http://127.0.0.1:18789\n"
    "Nhan Ctrl+C de dung"
)

para("Mở trình duyệt và truy cập: http://127.0.0.1:18789")

heading("3.8 Kiểm tra hoạt động", 2)
numbered("Mở trình duyệt tại http://127.0.0.1:18789 — xem trang Control UI")
numbered("Chạy lệnh kiểm tra sức khỏe:")
code_block("openclaw doctor")
numbered("Gửi tin nhắn thử qua Control UI hoặc channel đã cấu hình")
numbered("Kiểm tra log tại thư mục logs/")

doc.add_page_break()

# ============================================================
# 4. FILE CAU HINH
# ============================================================
heading("4. HƯỚNG DẪN FILE CẤU HÌNH", 1)

para(
    "Tất cả cấu hình người dùng nằm trong thư mục config/. "
    "Chỉ cần chỉnh các file này — KHÔNG sửa source code."
)

heading("4.1 config/llm_config.json — Cấu hình mô hình AI", 2)
tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = 'Table Grid'
for i, t in enumerate(["Trường", "Kiểu", "Bắt buộc", "Mô tả"]):
    tbl3.rows[0].cells[i].text = t
    tbl3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rows3 = [
    ("provider", "string", "Có", "Nhà cung cấp: anthropic | openai | google | xai | mistral | groq"),
    ("model", "string", "Có", "Tên model: claude-sonnet-4-6, gpt-4o, gemini-2.0-flash, ..."),
    ("api_key", "string", "Có", "API Key từ nhà cung cấp — KHÔNG chia sẻ key này"),
    ("base_url", "string", "Không", "URL tùy chỉnh cho model tự host (Ollama/LM Studio)"),
    ("fallback_provider", "string", "Không", "Provider dự phòng khi provider chính lỗi"),
    ("fallback_model", "string", "Không", "Model dự phòng"),
    ("fallback_api_key", "string", "Không", "API key dự phòng"),
    ("max_tokens", "number", "Không", "Số token tối đa mỗi lần trả lời (mặc định: 8192)"),
    ("temperature", "number", "Không", "Độ sáng tạo 0.0-2.0 (mặc định: 1.0)"),
]
for r in rows3:
    table_row(tbl3, r)
doc.add_paragraph()

heading("Ví dụ với Anthropic Claude:", 2)
code_block(
    '{\n'
    '  "provider": "anthropic",\n'
    '  "model": "claude-opus-4-8",\n'
    '  "api_key": "sk-ant-api03-xxxxx",\n'
    '  "base_url": ""\n'
    '}'
)

heading("Ví dụ với model tự host Ollama:", 2)
code_block(
    '{\n'
    '  "provider": "openai",\n'
    '  "model": "llama3.2",\n'
    '  "api_key": "ollama",\n'
    '  "base_url": "http://localhost:11434/v1"\n'
    '}'
)

heading("4.2 config/bot_config.json — Hành vi bot", 2)
tbl4 = doc.add_table(rows=1, cols=4)
tbl4.style = 'Table Grid'
for i, t in enumerate(["Trường", "Kiểu", "Bắt buộc", "Mô tả"]):
    tbl4.rows[0].cells[i].text = t
    tbl4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rows4 = [
    ("bot_name", "string", "Không", "Tên hiển thị của bot (mặc định: OpenClaw)"),
    ("language", "string", "Không", "Ngôn ngữ: vi (tiếng Việt), en (tiếng Anh)"),
    ("dm_policy", "string", "Không", "Chính sách DM: pairing | allowlist | open | disabled"),
    ("group_policy", "string", "Không", "Chính sách nhóm: disabled | open | allowlist"),
    ("allow_from", "array", "Không", "Danh sách ID được phép nhắn tin ([] = tất cả nếu open)"),
    ("workspace", "string", "Không", "Thư mục workspace (để trống = mặc định)"),
    ("agent_id", "string", "Không", "ID agent chính (mặc định: main)"),
]
for r in rows4:
    table_row(tbl4, r)
doc.add_paragraph()

para("Giải thích dm_policy:")
bullet("pairing: Người lạ nhận mã xác thực, cần phê duyệt thủ công (an toàn nhất)", level=0)
bullet("allowlist: Chỉ nhận tin từ các ID trong allow_from", level=0)
bullet("open: Nhận tin từ tất cả mọi người (cần thêm '*' vào allow_from)", level=0)
bullet("disabled: Tắt hoàn toàn", level=0)

heading("4.3 config/app_config.json — Cài đặt ứng dụng", 2)
tbl5 = doc.add_table(rows=1, cols=4)
tbl5.style = 'Table Grid'
for i, t in enumerate(["Trường", "Kiểu", "Mặc định", "Mô tả"]):
    tbl5.rows[0].cells[i].text = t
    tbl5.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rows5 = [
    ("gateway_port", "number", "18789", "Cổng HTTP của Gateway"),
    ("gateway_host", "string", "127.0.0.1", "Địa chỉ lắng nghe (0.0.0.0 = cho phép mạng LAN)"),
    ("auto_install_daemon", "boolean", "false", "Tự động cài service chạy ngầm"),
    ("verbose", "boolean", "false", "Hiển thị thông tin chi tiết khi chạy"),
    ("log_level", "string", "info", "Mức log: debug | info | warn | error"),
    ("control_ui_enabled", "boolean", "true", "Bật giao diện web quản lý"),
    ("sandbox_mode", "string", "off", "Chế độ sandbox: off | non-main | all"),
    ("remote_access", "boolean", "false", "Cho phép truy cập từ xa qua internet"),
]
for r in rows5:
    table_row(tbl5, r)
doc.add_paragraph()

heading("4.4 config/paths.json — Đường dẫn tùy chỉnh", 2)
tbl6 = doc.add_table(rows=1, cols=3)
tbl6.style = 'Table Grid'
for i, t in enumerate(["Trường", "Mặc định", "Mô tả"]):
    tbl6.rows[0].cells[i].text = t
    tbl6.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rows6 = [
    ("openclaw_config_dir", "%USERPROFILE%\\.openclaw", "Thư mục cấu hình OpenClaw"),
    ("workspace_dir", "%USERPROFILE%\\.openclaw\\workspace", "Thư mục workspace agent"),
    ("logs_dir", "logs", "Thư mục log cục bộ (tương đối với thư mục dự án)"),
    ("data_dir", "data", "Thư mục dữ liệu người dùng"),
    ("backups_dir", "backups", "Thư mục lưu backup"),
]
for r in rows6:
    table_row(tbl6, r)
doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 5. GIAI THICH LENH
# ============================================================
heading("5. GIẢI THÍCH CÁC LỆNH SỬ DỤNG", 1)

cmds = [
    ("install.bat", "Cài đặt toàn bộ hệ thống tự động — chạy một lần khi cài mới"),
    ("start.bat", "Khởi động OpenClaw Gateway — chạy mỗi khi muốn dùng"),
    ("stop.bat", "Dừng OpenClaw Gateway an toàn"),
    ("backup.bat", "Sao lưu toàn bộ cấu hình và dữ liệu"),
    ("restore.bat", "Khôi phục từ bản backup đã lưu"),
]

for cmd, desc in cmds:
    p = doc.add_paragraph()
    run = p.add_run(f"► {cmd}")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    para(f"   {desc}")

doc.add_paragraph()

heading("Lệnh openclaw (CLI):", 2)
cli_cmds = [
    ("openclaw --version", "Kiểm tra phiên bản đã cài"),
    ("openclaw doctor", "Kiểm tra sức khỏe hệ thống và phát hiện lỗi cấu hình"),
    ("openclaw doctor --fix", "Tự động sửa các lỗi phát hiện"),
    ("openclaw gateway --port 18789", "Khởi động Gateway tại cổng 18789"),
    ("openclaw gateway --port 18789 --verbose", "Khởi động với chế độ verbose (thông tin chi tiết)"),
    ("openclaw onboard", "Chạy trình hướng dẫn cài đặt tương tác"),
    ("openclaw onboard --install-daemon", "Cài OpenClaw như service chạy ngầm"),
    ("openclaw configure", "Mở trình cấu hình tương tác"),
    ("openclaw config get agents.defaults.model", "Xem giá trị một thuộc tính cấu hình"),
    ("openclaw config set agents.defaults.model.primary 'openai/gpt-4o'", "Đặt model chính"),
    ("openclaw update", "Cập nhật OpenClaw lên phiên bản mới nhất"),
    ("openclaw logs", "Xem log của Gateway"),
    ("openclaw status", "Kiểm tra trạng thái Gateway"),
    ("openclaw pairing approve telegram <code>", "Phê duyệt yêu cầu kết nối từ người dùng mới"),
    ("openclaw models list", "Liệt kê các model được hỗ trợ"),
]

tbl7 = doc.add_table(rows=1, cols=2)
tbl7.style = 'Table Grid'
for i, t in enumerate(["Lệnh", "Mô tả"]):
    tbl7.rows[0].cells[i].text = t
    tbl7.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for cmd, desc in cli_cmds:
    table_row(tbl7, [cmd, desc])
doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 6. LOI THUONG GAP
# ============================================================
heading("6. LỖI THƯỜNG GẶP VÀ CÁCH KHẮC PHỤC", 1)

errors = [
    {
        "error": "Lỗi: 'openclaw' is not recognized",
        "cause": "Node.js chưa được thêm vào PATH, hoặc cài đặt chưa hoàn tất",
        "fix": [
            "Đóng và mở lại cửa sổ Command Prompt",
            "Hoặc khởi động lại máy tính",
            "Kiểm tra: node --version hoạt động không",
            "Nếu vẫn lỗi, thêm thủ công vào PATH: C:\\Users\\<user>\\AppData\\Roaming\\npm",
        ],
    },
    {
        "error": "Lỗi: API key không hợp lệ (401 Unauthorized)",
        "cause": "API key sai hoặc chưa được điền vào config/llm_config.json",
        "fix": [
            "Mở config/llm_config.json",
            "Kiểm tra api_key đã điền đúng chưa (không có khoảng trắng thừa)",
            "Kiểm tra provider đúng với api_key (anthropic key cho OpenAI sẽ bị từ chối)",
            "Tạo API key mới tại trang của nhà cung cấp",
        ],
    },
    {
        "error": "Lỗi: Port 18789 đã được sử dụng",
        "cause": "Một ứng dụng khác hoặc instance OpenClaw cũ đang dùng cổng này",
        "fix": [
            "Chạy stop.bat để dừng instance cũ",
            "Hoặc thay đổi gateway_port trong config/app_config.json",
            "Kiểm tra cổng: netstat -ano | findstr :18789",
        ],
    },
    {
        "error": "Lỗi: npm install -g thất bại (EACCES / permission denied)",
        "cause": "Thiếu quyền ghi vào thư mục npm global",
        "fix": [
            "Chạy Command Prompt với quyền Administrator",
            "Hoặc đổi thư mục npm: npm config set prefix %APPDATA%\\npm",
        ],
    },
    {
        "error": "Lỗi: openclaw.json fails validation",
        "cause": "File cấu hình có trường không hợp lệ",
        "fix": [
            "Chạy: openclaw doctor để xem chi tiết lỗi",
            "Chạy: openclaw doctor --fix để tự sửa",
            "Hoặc xóa file ~/.openclaw/openclaw.json và chạy lại install.bat",
        ],
    },
    {
        "error": "Lỗi: Node.js version quá cũ",
        "cause": "Đang dùng Node.js < 22.14",
        "fix": [
            "Tải Node.js 24 LTS từ https://nodejs.org",
            "Gỡ cài Node.js cũ qua Control Panel trước",
            "Cài Node.js mới và chạy lại install.bat",
        ],
    },
]

for e in errors:
    p = doc.add_paragraph()
    run = p.add_run(f"❌ {e['error']}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p2 = doc.add_paragraph()
    p2.add_run("Nguyên nhân: ").bold = True
    p2.add_run(e["cause"])

    p3 = doc.add_paragraph()
    p3.add_run("Cách khắc phục:").bold = True
    for fix in e["fix"]:
        bullet(f"  {fix}", level=1)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 7. CAP NHAT
# ============================================================
heading("7. HƯỚNG DẪN CẬP NHẬT PHIÊN BẢN", 1)

heading("Cập nhật tự động (khuyến nghị):", 2)
code_block(
    "# Cập nhật lên phiên bản mới nhất\n"
    "npm install -g openclaw@latest\n\n"
    "# Hoặc dùng lệnh cập nhật tích hợp\n"
    "openclaw update\n\n"
    "# Sau khi cập nhật, kiểm tra\n"
    "openclaw --version\n"
    "openclaw doctor"
)

heading("Các kênh cập nhật:", 2)
bullet("stable (mặc định): Phiên bản ổn định — npm install -g openclaw@latest")
bullet("beta: Phiên bản thử nghiệm — npm install -g openclaw@beta")
bullet("dev: Phiên bản phát triển — npm install -g openclaw@dev")

heading("Trước khi cập nhật:", 2)
bullet("Chạy backup.bat để lưu cấu hình hiện tại")
bullet("Đọc CHANGELOG.md để biết thay đổi")
bullet("Chạy stop.bat để dừng Gateway")

doc.add_page_break()

# ============================================================
# 8. BACKUP VA RESTORE
# ============================================================
heading("8. BACKUP VÀ RESTORE CẤU HÌNH", 1)

heading("8.1 Tạo backup thủ công", 2)
code_block("# Nhấp đúp backup.bat hoặc chạy:\nbackup.bat")
para("Script sẽ tự động:")
bullet("Tạo thư mục backups/backup_YYYYMMDD_HHMM/")
bullet("Copy config/, data/, và ~/.openclaw/ vào đó")
bullet("Nén thành file .tar.gz")
bullet("Giữ lại 10 bản backup gần nhất, xóa bản cũ hơn")

heading("8.2 Lịch backup tự động (tùy chọn)", 2)
para("Thêm backup.bat vào Windows Task Scheduler để backup hàng ngày:")
numbered("Mở Task Scheduler (taskschd.msc)")
numbered("Create Task → Action → Start a program")
numbered("Program: C:\\path\\to\\openclaw\\backup.bat")
numbered("Trigger: Hàng ngày lúc 2:00 AM")

heading("8.3 Khôi phục từ backup", 2)
code_block("# Nhấp đúp restore.bat hoặc chạy:\nrestore.bat")
para("Script sẽ:")
bullet("Liệt kê tất cả bản backup có sẵn")
bullet("Cho phép bạn chọn bản muốn khôi phục")
bullet("Tự động backup cấu hình hiện tại trước khi ghi đè")
bullet("Khôi phục config/, data/, và ~/.openclaw/")

doc.add_page_break()

# ============================================================
# 9. SO DO HOAT DONG
# ============================================================
heading("9. SƠ ĐỒ HOẠT ĐỘNG HỆ THỐNG", 1)

para("Luồng hoạt động của OpenClaw:")
doc.add_paragraph()

flow_text = (
    "┌─────────────────────────────────────────────────────────────┐\n"
    "│                     LUỒNG HOẠT ĐỘNG                        │\n"
    "│                                                             │\n"
    "│  1. Người dùng chạy install.bat                            │\n"
    "│     └─► Kiểm tra Node.js, Git                              │\n"
    "│     └─► Cài openclaw từ npm                                │\n"
    "│     └─► Tạo thư mục (logs/, data/, backups/)              │\n"
    "│     └─► Apply config/ → ~/.openclaw/openclaw.json          │\n"
    "│                                                             │\n"
    "│  2. Người dùng chỉnh config/                               │\n"
    "│     ├── llm_config.json → API key, provider, model         │\n"
    "│     ├── bot_config.json → Tên bot, ngôn ngữ, policy        │\n"
    "│     ├── app_config.json → Cổng, log level                  │\n"
    "│     └── paths.json → Đường dẫn tùy chỉnh                  │\n"
    "│                                                             │\n"
    "│  3. Người dùng chạy start.bat                              │\n"
    "│     └─► Apply config → openclaw.json                       │\n"
    "│     └─► Đặt biến môi trường API key                        │\n"
    "│     └─► openclaw gateway --port 18789                      │\n"
    "│                                                             │\n"
    "│  4. OpenClaw Gateway chạy tại :18789                       │\n"
    "│     ├── Control UI: http://127.0.0.1:18789                 │\n"
    "│     ├── Nhận tin nhắn từ các kênh chat                     │\n"
    "│     ├── Gửi yêu cầu đến LLM API                           │\n"
    "│     └── Trả về phản hồi cho người dùng                     │\n"
    "│                                                             │\n"
    "│  5. Để dừng: Ctrl+C hoặc stop.bat                          │\n"
    "│  6. Backup định kỳ: backup.bat                             │\n"
    "└─────────────────────────────────────────────────────────────┘"
)
p = doc.add_paragraph()
run = p.add_run(flow_text)
run.font.name = "Courier New"
run.font.size = Pt(8)

doc.add_paragraph()

heading("Cây thư mục dự án:", 2)
tree_text = (
    "openclaw/\n"
    "├── config/                    ← Cấu hình người dùng\n"
    "│   ├── llm_config.json        ← API key, model AI\n"
    "│   ├── bot_config.json        ← Hành vi bot\n"
    "│   ├── app_config.json        ← Cài đặt ứng dụng\n"
    "│   ├── paths.json             ← Đường dẫn\n"
    "│   └── README.md              ← Hướng dẫn cấu hình\n"
    "├── logs/                      ← Log tự động theo ngày\n"
    "├── data/                      ← Dữ liệu người dùng\n"
    "├── backups/                   ← Backup tự động\n"
    "├── huongdan/                  ← Tài liệu hướng dẫn\n"
    "│   └── HuongDan_CaiDat_OpenClaw.docx\n"
    "├── scripts/                   ← Script hỗ trợ\n"
    "│   └── apply_openclaw_config.ps1\n"
    "├── install.bat                ← Cài đặt tự động\n"
    "├── start.bat                  ← Khởi động\n"
    "├── stop.bat                   ← Dừng\n"
    "├── backup.bat                 ← Sao lưu\n"
    "└── restore.bat                ← Khôi phục"
)
p2 = doc.add_paragraph()
run2 = p2.add_run(tree_text)
run2.font.name = "Courier New"
run2.font.size = Pt(9)

doc.add_paragraph()
separator()
para(
    "Tài liệu này được tạo tự động cho OpenClaw v2026.4.24. "
    "Xem thêm tại: https://docs.openclaw.ai",
    color=(0x7F, 0x8C, 0x8D),
)

# --- Luu file ---
doc.save(OUTPUT_PATH)
print(f"Da tao thanh cong: {OUTPUT_PATH}")
